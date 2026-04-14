"""
NEPOOL Summarizer -- scripts/summarize.py

Reads scraped_materials.json, downloads PDFs, extracts text with pdftotext,
calls Claude API to generate summaries, and outputs data/summaries.js.

Usage:
    python scripts/summarize.py                     # last 2 months, all committees
    python scripts/summarize.py --months 1          # just last month
    python scripts/summarize.py --committee mc      # one committee only
    python scripts/summarize.py --dry-run           # show what would run, no API calls
    python scripts/summarize.py --upcoming-only     # only docs for upcoming meetings
                                                    # that match a real agenda item

Output: data/summaries.js  (window.SUMMARIES_DATA global -- loaded by index.html)
"""

import anthropic
import json
import os
import re
import shutil
import subprocess
import sys
import urllib.request
from datetime import date, timedelta
from pathlib import Path

MEETINGS_FILE = Path("data/meetings.js")

SCRAPED_FILE  = Path("data/scraped_materials.json")
SUMMARIES_FILE = Path("data/summaries.js")


# ── Scraped-data loader (json preferred; js fallback) ─────────────────────────

def load_scraped_data():
    """
    Load scraped_materials into a dict.
    Prefers the .json file (written by scrape_materials.py at runtime).
    Falls back to the committed .js file so the script can run locally
    without needing to re-scrape first.
    """
    if SCRAPED_FILE.exists():
        with open(SCRAPED_FILE, encoding="utf-8") as f:
            return json.load(f)
    js_path = SCRAPED_FILE.with_suffix(".js")
    if js_path.exists():
        raw = js_path.read_text(encoding="utf-8")
        m = re.search(r'window\.\w+\s*=\s*', raw)
        if m:
            return json.loads(raw[m.end():].rstrip().rstrip(';').rstrip())
    raise FileNotFoundError(
        f"Neither {SCRAPED_FILE} nor {SCRAPED_FILE.with_suffix('.js')} found. "
        "Run scrape_materials.py first."
    )
PDF_CACHE_DIR  = Path("data/pdfs")

# Max characters sent to Claude per document (controls cost and avoids token limits)
MAX_TEXT_CHARS = 12_000

# File types we cannot summarize — skip them
SKIP_EXTENSIONS = {"xlsx", "xls", "doc", "zip"}

# Skip documents whose titles match these patterns (procedural boilerplate)
SKIP_TITLE_PATTERNS = [
    r"\bbylaw",
    r"\battendance\b",
    r"guest.*form",
    r"two.month.*look.ahead",
    r"annual.*work.*plan",
]

# ── Claude prompt ─────────────────────────────────────────────────────────────

SYSTEM_PROMPT = """\
You are summarizing NEPOOL committee documents for a Maine PUC analyst who tracks \
ISO-NE wholesale market issues, with particular focus on impacts to Maine ratepayers \
and Maine's utilities (CMP, Versant).

Respond ONLY with valid JSON — no markdown fences, no extra text. Use this exact structure:
{
  "summary": "2–3 sentence plain-English summary of what is presented or decided",
  "item_type": "Presentation" | "Vote" | "Informational" | "Procedural" | "Minutes" | "Actions",
  "maine_relevance": "high" | "medium" | "low",
  "maine_relevance_reason": "One sentence explaining the relevance rating",
  "maine_mentions": ["explicit mentions of Maine, CMP, Versant, or Maine-specific data"],
  "topic_tags": ["short-kebab-case", "topic", "tags"]
}

Common topic tags: capacity-markets, energy-markets, transmission-planning, cost-allocation, \
clean-energy, load-forecasting, demand-response, ancillary-services, interconnection, \
tariff, compliance, regional-planning, maine-specific, forecast, settlement, metering."""


# ── Helpers ───────────────────────────────────────────────────────────────────

def should_skip(doc):
    """Return True if this document should not be summarized."""
    if doc["ext"] in SKIP_EXTENSIONS:
        return True
    title_lower = doc["title"].lower()
    return any(re.search(p, title_lower) for p in SKIP_TITLE_PATTERNS)


def safe_filename(url):
    """Derive a safe local filename from a URL."""
    name = url.split("/")[-1].split("?")[0]
    name = re.sub(r'[<>:"|?*\\]', "_", name)
    return name or "document.pdf"


def find_pdftotext():
    """Locate pdftotext — check known path first, then PATH."""
    known = Path(
        r"C:\Users\micha\AppData\Local\Microsoft\WinGet\Packages"
        r"\oschwartz10612.Poppler_Microsoft.Winget.Source_8wekyb3d8bbwe"
        r"\poppler-25.07.0\Library\bin\pdftotext.exe"
    )
    if known.exists():
        return str(known)
    found = shutil.which("pdftotext")
    if found:
        return found
    return None


def download_pdf(url, dest_path):
    """Download a file to dest_path if not already cached. Returns True on success."""
    if dest_path.exists():
        return True
    try:
        req = urllib.request.Request(
            url, headers={"User-Agent": "Mozilla/5.0 (compatible; NEPOOL-Tracker/1.0)"}
        )
        with urllib.request.urlopen(req, timeout=30) as resp:
            dest_path.write_bytes(resp.read())
        return True
    except Exception as e:
        print(f"      Download failed: {e}")
        return False


def extract_text(file_path, pdftotext_cmd):
    """Extract plain text from a PDF or DOCX. Returns string or None."""
    if str(file_path).lower().endswith(".docx"):
        return extract_text_docx(file_path)
    return extract_text_pdf(file_path, pdftotext_cmd)


def extract_text_pdf(pdf_path, pdftotext_cmd):
    try:
        result = subprocess.run(
            [pdftotext_cmd, "-layout", str(pdf_path), "-"],
            capture_output=True,
            encoding="utf-8",
            errors="replace",   # swap undecodable bytes with ? rather than crashing
            timeout=30
        )
        if result.returncode == 0:
            return result.stdout.strip()
        print(f"      pdftotext error (rc={result.returncode}): {result.stderr[:200]}")
        return None
    except Exception as e:
        print(f"      Text extraction failed: {e}")
        return None


def extract_text_docx(docx_path):
    """
    Extract text from a DOCX, including both paragraphs and table cells.
    Many ISO-NE agendas are formatted as tables, so table extraction is critical.
    Merged cells repeat their text across columns — we deduplicate within each row.
    """
    try:
        import docx as docx_lib
        doc = docx_lib.Document(str(docx_path))
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                seen, unique = set(), []
                for cell in row.cells:
                    t = cell.text.strip()
                    if t and t not in seen:
                        seen.add(t)
                        unique.append(t)
                if unique:
                    lines.append(" | ".join(unique))
        return "\n".join(lines)
    except Exception as e:
        print(f"      DOCX text extraction failed: {e}")
        return None


def call_claude(client, title, committee_name, meeting_date, text):
    """
    Call Claude API and return a parsed summary dict.
    Returns None on failure.

    We use a structured JSON prompt so the output can be stored and displayed
    without any further parsing — this is a best practice for API integrations
    that need reliable, machine-readable output.
    """
    truncated = text[:MAX_TEXT_CHARS]
    if len(text) > MAX_TEXT_CHARS:
        truncated += "\n\n[... document truncated for length ...]"

    user_msg = (
        f"Committee: {committee_name}\n"
        f"Meeting date: {meeting_date}\n"
        f"Document title: {title}\n\n"
        f"Document text:\n{truncated}"
    )

    try:
        response = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=600,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": user_msg}],
        )
        raw = response.content[0].text.strip()
        return json.loads(raw)
    except json.JSONDecodeError:
        print(f"      Claude returned invalid JSON — skipping.")
        print(f"      Raw: {raw[:300]}")
        return None
    except Exception as e:
        print(f"      Claude API error: {e}")
        return None


def load_existing_summaries():
    """
    Load previously generated summaries.js so we can skip docs already done.
    Returns the inner dict (committees → dates → list), or {} if none.
    """
    if not SUMMARIES_FILE.exists():
        return {}
    js_text = SUMMARIES_FILE.read_text(encoding="utf-8")
    m = re.search(r"window\.SUMMARIES_DATA\s*=\s*(\{.*\})\s*;", js_text, re.DOTALL)
    if not m:
        return {}
    try:
        data = json.loads(m.group(1))
        return data.get("committees", {})
    except json.JSONDecodeError:
        print("Warning: could not parse existing summaries.js — starting fresh.")
        return {}


# ── Upcoming-only helpers ──────────────────────────────────────────────────────

def load_upcoming_agenda_index():
    """
    Parse meetings.js and return a dict mapping (committee_id, meeting_date)
    to the list of agenda_numbers for all ACTIVE meetings that have real
    agenda items (i.e., at least one item has an agenda_number).

    'Active' means: upcoming OR occurred within the last 14 days. The 14-day
    lookback lets post-meeting materials (Composite, NOA, Actions Letter) get
    summarized after the meeting without needing a separate pipeline step.
    """
    if not MEETINGS_FILE.exists():
        return {}

    raw = MEETINGS_FILE.read_text(encoding="utf-8")
    raw = re.sub(r'(?m)^\s*//[^\n]*\n?', '', raw)
    m = re.search(r'window\.\w+\s*=\s*', raw)
    if not m:
        return {}
    raw = raw[m.end():].rstrip().rstrip(';').rstrip()
    raw = re.sub(r'(?m)^(\s*)([a-zA-Z_$][a-zA-Z0-9_$]*)(\s*):', r'\1"\2"\3:', raw)
    raw = re.sub(r',(\s*[}\]])', r'\1', raw)

    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        return {}

    today = date.today()
    lookback = today - timedelta(days=14)
    index = {}  # (cid, meeting_date) -> [agenda_number, ...]

    for committee in data.get("committees", []):
        cid = committee.get("id", "")
        for meeting in committee.get("meetings", []):
            mdate = meeting.get("date", "")
            if not mdate:
                continue
            try:
                if date.fromisoformat(mdate) < lookback:
                    continue
            except ValueError:
                continue
            numbers = [
                item["agenda_number"]
                for item in meeting.get("agenda_items", [])
                if item.get("agenda_number")
            ]
            if numbers:
                index[(cid, mdate)] = numbers

    return index


def matches_any_agenda_number(title, agenda_numbers):
    """
    Return True if the doc title matches at least one agenda_number using
    the same word-boundary regex as index.html's getScrapedDocs().
    """
    for num in agenda_numbers:
        esc = re.escape(num)
        if re.search(r'\b' + esc + r'([ .\-(]|$)', title, re.IGNORECASE):
            return True
    return False


# ── Main ──────────────────────────────────────────────────────────────────────

def main(months_back, committee_filter, dry_run, upcoming_only=False):
    try:
        scraped = load_scraped_data()
    except FileNotFoundError as e:
        print(f"ERROR: {e}")
        sys.exit(1)

    pdftotext_cmd = find_pdftotext()
    if not pdftotext_cmd:
        print("ERROR: pdftotext not found. Install poppler or add it to PATH.")
        sys.exit(1)

    cutoff = date.today() - timedelta(days=months_back * 30)

    print("NEPOOL Summarizer")
    print(f"Cutoff date:  {cutoff}  ({months_back} months back)")
    print(f"Committee:    {committee_filter or 'all'}")
    print(f"Dry run:      {dry_run}")
    print(f"Upcoming only: {upcoming_only}")
    print(f"pdftotext:    {pdftotext_cmd}\n")

    # Load agenda index for --upcoming-only filtering
    agenda_index = load_upcoming_agenda_index() if upcoming_only else {}
    if upcoming_only:
        if agenda_index:
            print(f"Active meetings with real agendas (upcoming + last 14 days): {len(agenda_index)}")
            for (cid, mdate), nums in sorted(agenda_index.items()):
                print(f"  {cid.upper()} {mdate}: {len(nums)} agenda item(s)")
        else:
            print("No active meetings with real agenda items found. Nothing to summarize.")
        print()

    # Set up Claude client (unless dry run)
    if not dry_run:
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            print("ERROR: ANTHROPIC_API_KEY environment variable not set.")
            print("In PowerShell: $env:ANTHROPIC_API_KEY = 'your-key-here'")
            sys.exit(1)
        client = anthropic.Anthropic(api_key=api_key)
    else:
        client = None

    PDF_CACHE_DIR.mkdir(parents=True, exist_ok=True)

    # Load existing summaries to avoid re-processing
    existing = load_existing_summaries()
    reused_count = sum(
        len(docs)
        for cdata in existing.values()
        for docs in cdata.values()
    )
    if reused_count:
        print(f"Found {reused_count} existing summaries — will reuse where URLs match.\n")

    # Build results, starting from existing so unprocessed committees are preserved
    results = {
        "generated_at": date.today().isoformat(),
        "committees": dict(existing),  # copy existing; overwrite committees we re-process
    }

    total_new = total_reused = total_skipped = 0

    for cid, cdata in scraped["committees"].items():
        if committee_filter and cid != committee_filter:
            continue

        committee_name = cdata["name"]
        meetings = cdata.get("meetings", {})

        print(f"[{cid.upper()}] {committee_name}")

        if cid not in results["committees"]:
            results["committees"][cid] = {}

        for meeting_date in sorted(meetings.keys(), reverse=True):
            if date.fromisoformat(meeting_date) < cutoff:
                continue

            # --upcoming-only: skip past meetings and meetings without real agendas
            if upcoming_only:
                agenda_numbers = agenda_index.get((cid, meeting_date))
                if not agenda_numbers:
                    continue  # not an upcoming meeting with a real agenda
            else:
                agenda_numbers = None

            docs = meetings[meeting_date]
            to_process = [d for d in docs if not should_skip(d)]

            # --upcoming-only: further filter to docs matching an agenda_number
            if agenda_numbers is not None:
                to_process = [
                    d for d in to_process
                    if matches_any_agenda_number(d["title"], agenda_numbers)
                ]

            skipped_type = len(docs) - len(to_process)

            print(f"  {meeting_date}: {len(to_process)} PDFs  ({skipped_type} non-PDF skipped)")

            # Existing summaries for this meeting, keyed by URL
            existing_by_url = {
                s["url"]: s
                for s in existing.get(cid, {}).get(meeting_date, [])
            }

            meeting_summaries = []

            for doc in to_process:
                url   = doc["url"]
                title = doc["title"]

                # Reuse if we already have a summary for this exact URL
                if url in existing_by_url:
                    meeting_summaries.append(existing_by_url[url])
                    total_reused += 1
                    continue

                if dry_run:
                    print(f"    [DRY RUN] {title[:72]}")
                    total_new += 1
                    continue

                print(f"    >>  {title[:72]}")

                # Download PDF to cache
                pdf_dir  = PDF_CACHE_DIR / cid / meeting_date
                pdf_dir.mkdir(parents=True, exist_ok=True)
                pdf_path = pdf_dir / safe_filename(url)

                if not download_pdf(url, pdf_path):
                    total_skipped += 1
                    continue

                # Extract text
                text = extract_text(pdf_path, pdftotext_cmd)
                if not text or len(text.strip()) < 50:
                    print("      Skipping — no usable text extracted.")
                    total_skipped += 1
                    continue

                # Summarize
                result = call_claude(client, title, committee_name, meeting_date, text)
                if result is None:
                    total_skipped += 1
                    continue

                entry = {"title": title, "url": url, "ext": doc["ext"], **result}
                meeting_summaries.append(entry)
                total_new += 1

                rel = result["maine_relevance"].upper()
                print(f"      [{rel}] {result['summary'][:80]}...")

            results["committees"][cid][meeting_date] = meeting_summaries

        print()

    # Write output
    if not dry_run:
        js_out = (
            "// Auto-generated by summarize.py — do not edit manually.\n"
            "// Re-run to refresh; existing summaries are reused (cached by URL).\n"
            f"window.SUMMARIES_DATA = {json.dumps(results, indent=2, ensure_ascii=False)};\n"
        )
        SUMMARIES_FILE.write_text(js_out, encoding="utf-8")
        print(f"Saved -> {SUMMARIES_FILE}")

    print(f"\nResults: {total_new} new  |  {total_reused} reused  |  {total_skipped} failed/skipped")
    if dry_run:
        print("(Dry run — no files written, no API calls made.)")


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    months = 2
    committee = None
    dry_run = False
    upcoming_only = False

    args = sys.argv[1:]
    i = 0
    while i < len(args):
        arg = args[i]
        if arg == "--months" and i + 1 < len(args):
            months = int(args[i + 1]); i += 2
        elif arg.startswith("--months="):
            months = int(arg.split("=", 1)[1]); i += 1
        elif arg == "--committee" and i + 1 < len(args):
            committee = args[i + 1].lower(); i += 2
        elif arg.startswith("--committee="):
            committee = arg.split("=", 1)[1].lower(); i += 1
        elif arg == "--dry-run":
            dry_run = True; i += 1
        elif arg == "--upcoming-only":
            upcoming_only = True; i += 1
        else:
            i += 1

    main(months, committee, dry_run, upcoming_only)
